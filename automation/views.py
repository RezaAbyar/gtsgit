from docx import Document
import os
from django.core.files.base import ContentFile
from django.urls import reverse_lazy
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User, Group
from django.core.paginator import Paginator
from django.db.models import Q
from django.template.response import TemplateResponse
from datetime import datetime
from base.models import Owner
from base.permission_decoder import cache_permission
from .models import Message, PngImage, MessageReadLog
from .forms import MessageForm
from PIL import Image
from io import BytesIO
import base64
from django.conf import settings
from django.views.generic import ListView, CreateView, UpdateView, DeleteView

Image.MAX_IMAGE_PIXELS = None  # غیرفعال کردن محدودیت اندازه تصویر


def display_tif(image_path):
    try:
        img = Image.open(image_path)
        return img
    except Exception as e:
        print(f"Error opening TIF file: {e}")
        return None


@cache_permission('message')
def inbox(request):
    messages = Message.objects.filter(
        Q(recipients=request.user.owner) | Q(groups=request.user.owner.role_id)
    ).distinct().order_by('-sent_at')
    messages = messages.filter(Q(zone_id=request.user.owner.zone_id) | Q(zone_id__isnull=True))
    messages = messages.filter(Q(area_id=request.user.owner.area_id) | Q(area_id__isnull=True))

    # Pagination
    paginator = Paginator(messages, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    return TemplateResponse(request, 'messages/inbox.html', {'page_obj': page_obj})


@cache_permission('message')
def sent_messages(request):
    messages = Message.objects.filter(sender=request.user.owner).order_by('-sent_at')

    # Pagination
    paginator = Paginator(messages, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    return TemplateResponse(request, 'messages/sent.html', {'page_obj': page_obj})


@cache_permission('message')
def compose(request, parent_id=None):
    parent_message = None
    recipients = Owner.object_role.c_owner(request).filter(Q(role__role__in=['gs', 'tek']) | Q(refrence_id=1))
    if parent_id:
        parent_message = get_object_or_404(Message, id=parent_id)

    if request.method == 'POST':
        form = MessageForm(request.POST, request.FILES, sender=request.user.owner)
        if form.is_valid():
            message = form.save(commit=False)
            message.sender = request.user.owner

            if request.user.owner.role.role == 'zone':
                message.zone = request.user.owner.zone
            if request.user.owner.role.role == 'area':
                message.area = request.user.owner.area

            # اگر گزینه "ذخیره پیش‌نویس" زده شده است
            if 'save_draft' in request.POST:
                message.is_draft = True
                message.save()
                form.save_m2m()
                return redirect('automation:inbox')

            # اگر قالب ورد انتخاب شده است، یک کپی از آن ایجاد کنید
            if message.template:
                template_path = message.template.template_file.path
                doc = Document(template_path)

                # ذخیره فایل ورد جدید با نام منحصر به فرد
                draft_filename = f"draft_{os.path.basename(template_path)}"
                output_path = os.path.join(settings.MEDIA_ROOT, 'draft_word_files', draft_filename)
                doc.save(output_path)

                # ذخیره فایل ورد در مدل Message
                with open(output_path, 'rb') as f:
                    message.word_file.save(draft_filename, ContentFile(f.read()))

            # ذخیره نهایی پیام
            message.save()
            form.save_m2m()

            # تبدیل فایل ورد به PNG (در صورت نیاز)
            if message.word_file and not message.is_draft:
                data = {
                    'user_name': request.user.owner.get_full_name(),
                    'date': datetime.now().strftime("%Y/%m/%d"),
                    'letter_number': "12345"  # شماره نامه می‌تواند از سیستم دریافت شود
                }
                filled_template_path = os.path.join(settings.MEDIA_ROOT, 'filled_word_files',
                                                    f"filled_{os.path.basename(message.word_file.path)}")
                fill_word_template(message.word_file.path, filled_template_path, data)

                # ذخیره فایل پر شده
                with open(filled_template_path, 'rb') as f:
                    message.word_file.save(f"filled_{os.path.basename(message.word_file.path)}", ContentFile(f.read()))
                print('png')
                # این بخش نیاز به پیاده‌سازی تبدیل DOCX به PNG دارد
                pass

            return redirect('automation:inbox')
    else:
        initial = {}
        if parent_message:
            initial = {
                'subject': f"Re: {parent_message.subject}",
                'recipients': [parent_message.sender.id],
            }
        form = MessageForm(sender=request.user, initial=initial)

    return TemplateResponse(request, 'messages/compose.html', {
        'form': form,
        'parent_message': parent_message,
        'recipients': recipients,
    })


def fill_word_template(template_path, output_path, data):
    doc = Document(template_path)

    # جایگزینی متغیرها در فایل ورد
    for paragraph in doc.paragraphs:
        if '{{نام کاربر}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{نام کاربر}}', data['user_name'])
        if '{{تاریخ}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{تاریخ}}', data['date'])
        if '{{شماره نامه}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{شماره نامه}}', data['letter_number'])

    doc.save(output_path)


@cache_permission('message')
def message_detail(request, message_id):
    message = get_object_or_404(Message, id=message_id)
    # ابتدا کوئری برای شرط اول

    # recipients = Owner.object_role.c_owner(request).filter(role__role__in=['gs','tek'])
    # Mark as read if recipient
    # if request.user.owner in message.recipients.all():

    message.mark_as_read(request.user)

    # Get PNG pages if exists
    png_pages = PngImage.objects.filter(tiff_image=message).count()
    replies = Message.objects.filter(parent_message=message).order_by('sent_at')

    return TemplateResponse(request, 'messages/detail.html', {
        'message': message,
        'replies': replies,
        'png_pages': png_pages,

    })


@cache_permission('message')
def image_viewer(request, message_id):
    message = get_object_or_404(Message, id=message_id)
    total_pages = PngImage.objects.filter(tiff_image=message).count()
    current_page = 1  # صفحه پیش‌فرض
    if request.method == 'POST':
        current_page = int(request.POST.get('page', 1))

    image_data = get_object_or_404(PngImage, tiff_image_id=message_id, page_number=current_page)

    return TemplateResponse(request, 'messages/image_viewer.html', {
        'message': message,
        'image_data': image_data,
        'current_page': current_page,
        'total_pages': total_pages,
    })


@cache_permission('message')
def read_reports(request, message_id):
    logs = MessageReadLog.objects.filter(message_id=message_id).order_by('-read_at')
    paginator = Paginator(logs, 100)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    return TemplateResponse(request, 'messages/read_reports.html', {'page_obj': page_obj})


from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from .models import WordTemplate


class WordTemplateListView(ListView):
    model = WordTemplate
    template_name = 'messages/word_template_list.html'


class WordTemplateCreateView(CreateView):
    model = WordTemplate
    fields = ['name', 'template_file']
    template_name = 'messages/word_template_form.html'
    success_url = reverse_lazy('automation:word_template_list')


class WordTemplateUpdateView(UpdateView):
    model = WordTemplate
    fields = ['name', 'template_file']
    template_name = 'messages/word_template_form.html'
    success_url = reverse_lazy('automation:word_template_list')


class WordTemplateDeleteView(DeleteView):
    model = WordTemplate
    template_name = 'messages/word_template_confirm_delete.html'
    success_url = reverse_lazy('automation:word_template_list')
